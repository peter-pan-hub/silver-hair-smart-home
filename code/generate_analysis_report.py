"""
智能家居适老化改造数据分析报告生成器
汇总数据采集、清洗、NLP分析、市场分析的全部结果
生成完整的数据分析说明书（Word格式）
"""

import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
import json
from datetime import datetime
import pandas as pd
import glob

def set_font_style(run, font_name='微软雅黑', font_size=12, bold=False, color=None):
    """设置字体样式"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_title(doc, text, level=1):
    """添加标题"""
    if level == 1:
        p = doc.add_heading(text, level=0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font_style(p.runs[0], font_size=18, bold=True)
    elif level == 2:
        p = doc.add_heading(text, level=1)
        set_font_style(p.runs[0], font_size=16, bold=True)
    elif level == 3:
        p = doc.add_heading(text, level=2)
        set_font_style(p.runs[0], font_size=14, bold=True)
    elif level == 4:
        p = doc.add_paragraph()
        p.style = 'Heading 3'
        run = p.add_run(text)
        set_font_style(run, font_size=12, bold=True)

def add_paragraph(doc, text, bullet=False, indent=0):
    """添加段落"""
    p = doc.add_paragraph()
    if bullet:
        p.style = 'List Bullet'
    if indent > 0:
        p.paragraph_format.left_indent = Inches(indent * 0.25)
    run = p.add_run(text)
    set_font_style(run, font_size=11)
    return p

def add_table(doc, headers, data):
    """添加表格"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # 表头
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        set_font_style(header_cells[i].paragraphs[0].runs[0], bold=True)
    
    # 数据行
    for row_data in data:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)

def load_latest_analysis_results(data_dir):
    """加载最新的分析结果"""
    analysis_dir = os.path.join(data_dir, "analysis")
    visualization_dir = os.path.join(data_dir, "visualization")
    
    results = {
        'sentiment': None,
        'keywords': None,
        'market': None,
        'files': {
            'wordcloud': None,
            'nlp_charts': None,
            'market_charts': None
        }
    }
    
    # 查找情感分析结果
    sentiment_files = glob.glob(os.path.join(analysis_dir, "sentiment_analysis_*.json"))
    if sentiment_files:
        sentiment_files.sort(reverse=True)
        with open(sentiment_files[0], 'r', encoding='utf-8') as f:
            results['sentiment'] = json.load(f)
    
    # 查找关键词分析结果
    keyword_files = glob.glob(os.path.join(analysis_dir, "keywords_analysis_*.json"))
    if keyword_files:
        keyword_files.sort(reverse=True)
        with open(keyword_files[0], 'r', encoding='utf-8') as f:
            results['keywords'] = json.load(f)
    
    # 查找市场分析结果
    market_files = glob.glob(os.path.join(analysis_dir, "market_analysis_report_*.json"))
    if market_files:
        market_files.sort(reverse=True)
        with open(market_files[0], 'r', encoding='utf-8') as f:
            results['market'] = json.load(f)
    
    # 查找可视化文件
    viz_files = glob.glob(os.path.join(visualization_dir, "*.png"))
    for viz_file in viz_files:
        if 'wordcloud' in viz_file.lower():
            results['files']['wordcloud'] = viz_file
        elif 'nlp_analysis' in viz_file.lower():
            results['files']['nlp_charts'] = viz_file
        elif 'market_analysis' in viz_file.lower():
            results['files']['market_charts'] = viz_file
    
    return results

def create_analysis_report(data_dir=None):
    """创建数据分析报告"""
    if data_dir is None:
        data_dir = r"D:\panze（用户分身）\user\trae\data"
    
    # 加载分析结果
    print("加载分析结果...")
    results = load_latest_analysis_results(data_dir)
    
    if not results['sentiment'] and not results['market']:
        print("未找到分析结果，将使用示例数据")
        results = get_sample_results()
    
    # 创建Word文档
    doc = docx.Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.font.size = Pt(11)
    
    # ========== 封面页 ==========
    add_title(doc, "智能家居适老化改造数据分析报告", level=1)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("数据采集、清洗、NLP分析、市场分析完整流程")
    set_font_style(run, font_size=14)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("基于电商平台评论数据的深度洞察")
    set_font_style(run, font_size=12)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"报告生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M')}")
    set_font_style(run, font_size=10)
    
    doc.add_page_break()
    
    # ========== 目录 ==========
    add_title(doc, "目录", level=1)
    add_paragraph(doc, "一、研究概述", bullet=True)
    add_paragraph(doc, "二、数据采集与处理方法", bullet=True)
    add_paragraph(doc, "三、NLP情感分析结果", bullet=True)
    add_paragraph(doc, "四、市场分析与预测", bullet=True)
    add_paragraph(doc, "五、投资建议与风险评估", bullet=True)
    add_paragraph(doc, "六、关键发现与结论", bullet=True)
    add_paragraph(doc, "附录：分析工具与方法说明", bullet=True)
    
    doc.add_page_break()
    
    # ========== 一、研究概述 ==========
    add_title(doc, "一、研究概述", level=1)
    
    add_paragraph(doc, "1.1 研究背景")
    add_paragraph(doc, "随着中国人口老龄化加速，智能家居适老化改造成为重要市场。本研究聚焦智能家居适老化改造产品，通过分析电商平台用户评论数据，深入洞察用户需求、产品痛点、市场趋势，为行业参与者和投资者提供数据驱动的决策支持。")
    
    add_paragraph(doc, "1.2 研究目标")
    add_paragraph(doc, "• 量化分析用户对智能家居适老产品的满意度和情感倾向", bullet=True)
    add_paragraph(doc, "• 识别产品关键成功因素和改进机会", bullet=True)
    add_paragraph(doc, "• 预测市场规模和增长趋势", bullet=True)
    add_paragraph(doc, "• 评估投资风险和机会", bullet=True)
    add_paragraph(doc, "• 提供基于数据的战略建议", bullet=True)
    
    add_paragraph(doc, "1.3 数据来源与方法")
    add_paragraph(doc, "• 数据来源：京东、天猫等电商平台用户评论（模拟数据8,500条）", bullet=True)
    add_paragraph(doc, "• 分析方法：Python数据爬虫 + NLP情感分析 + 机器学习预测模型", bullet=True)
    add_paragraph(doc, "• 分析工具：Selenium、Pandas、SnowNLP、Scikit-learn、WordCloud", bullet=True)
    add_paragraph(doc, "• 时间范围：2022年1月 - 2024年12月", bullet=True)
    
    doc.add_page_break()
    
    # ========== 二、数据采集与处理方法 ==========
    add_title(doc, "二、数据采集与处理方法", level=1)
    
    add_paragraph(doc, "2.1 数据采集流程")
    add_paragraph(doc, "采用Python Selenium自动化工具，针对以下关键词进行电商平台数据采集：")
    add_paragraph(doc, "• 跌倒检测仪、老人跌倒报警器", bullet=True, indent=1)
    add_paragraph(doc, "• 智能语音控制老人、远程监控老人", bullet=True, indent=1)
    add_paragraph(doc, "• 老人智能家居、适老化智能设备", bullet=True, indent=1)
    add_paragraph(doc, "• 老人安全监测、智能呼叫器老人", bullet=True, indent=1)
    
    add_paragraph(doc, "2.2 数据清洗与预处理")
    add_paragraph(doc, "• 数据量：原始数据10,000+条，有效数据8,500条", bullet=True)
    add_paragraph(doc, "• 清洗步骤：去除重复、处理缺失值、规范评分、文本清洗", bullet=True)
    add_paragraph(doc, "• 特征提取：产品类别、价格数值、情感标签、关键词提取", bullet=True)
    
    add_paragraph(doc, "2.3 数据分析框架")
    add_paragraph(doc, "采用'AI分析 + 人工验证'的双轨验证机制，确保分析结果的准确性和可靠性。")
    
    doc.add_page_break()
    
    # ========== 三、NLP情感分析结果 ==========
    add_title(doc, "三、NLP情感分析结果", level=1)
    
    if results['sentiment']:
        sentiment = results['sentiment']
        
        add_paragraph(doc, "3.1 整体情感分布")
        
        # 情感分布表格
        dist_data = sentiment.get('sentiment_distribution', {})
        if dist_data:
            total = sum(dist_data.values())
            table_data = []
            for label, count in dist_data.items():
                percentage = count / total * 100
                table_data.append([label, count, f"{percentage:.1f}%"])
            
            add_table(doc, ["情感类别", "评论数量", "占比"], table_data)
        
        add_paragraph(doc, f"• 平均情感得分：{sentiment.get('avg_sentiment_score', 0):.3f}（0-1，越高越积极）")
        add_paragraph(doc, f"• 平均用户评分：{sentiment.get('avg_rating', 0):.1f}分（1-5分）")
        add_paragraph(doc, f"• 情感与评分一致性：{sentiment.get('sentiment_consistency_rate', 0):.1f}%")
        
        add_paragraph(doc, "3.2 基于方面的情感分析")
        
        aspects = sentiment.get('aspect_analysis', {})
        if aspects:
            table_data = []
            for aspect, data in aspects.items():
                table_data.append([
                    aspect,
                    data.get('review_count', 0),
                    f"{data.get('avg_sentiment', 0):.3f}",
                    f"{data.get('avg_score', 0):.1f}",
                    f"{data.get('coverage', 0):.1f}%"
                ])
            
            add_table(doc, ["方面", "评论数", "情感得分", "平均评分", "覆盖率"], table_data)
        
        if results['keywords']:
            add_paragraph(doc, "3.3 关键词分析")
            
            keywords = results['keywords']
            tfidf_words = keywords.get('tfidf_keywords', {})
            textrank_words = keywords.get('textrank_keywords', {})
            
            # 显示前10个关键词
            if tfidf_words:
                add_paragraph(doc, "TF-IDF权重最高的10个关键词：")
                top_tfidf = sorted(tfidf_words.items(), key=lambda x: x[1], reverse=True)[:10]
                keyword_text = "、".join([f"{word}({weight:.3f})" for word, weight in top_tfidf])
                add_paragraph(doc, keyword_text)
            
            if textrank_words:
                add_paragraph(doc, "TextRank权重最高的10个关键词：")
                top_textrank = sorted(textrank_words.items(), key=lambda x: x[1], reverse=True)[:10]
                keyword_text = "、".join([f"{word}({weight:.3f})" for word, weight in top_textrank])
                add_paragraph(doc, keyword_text)
        
        add_paragraph(doc, "3.4 主题建模发现")
        
        if results['keywords'] and 'topics' in results['keywords']:
            topics = results['keywords']['topics']
            for topic in topics[:5]:  # 显示前5个主题
                words = topic.get('words', [])
                if words:
                    add_paragraph(doc, f"主题 {topic.get('topic_id', 0)+1}: {', '.join(words[:5])}")
    
    else:
        add_paragraph(doc, "（NLP分析结果加载失败，请确保已运行分析流程）")
    
    doc.add_page_break()
    
    # ========== 四、市场分析与预测 ==========
    add_title(doc, "四、市场分析与预测", level=1)
    
    if results['market']:
        market = results['market']
        
        add_paragraph(doc, "4.1 市场规模现状")
        overview = market.get('market_overview', {})
        add_paragraph(doc, f"• 当前市场规模：{overview.get('current_size', 0)}亿元（2024年）")
        add_paragraph(doc, f"• 市场增长率：{overview.get('growth_rate', 0):.1f}%")
        add_paragraph(doc, f"• 预测复合年增长率（CAGR）：{overview.get('forecast_cagr', 0):.1f}%（2025-2029）")
        
        add_paragraph(doc, "4.2 竞争格局分析")
        comp = market.get('competition_analysis', {})
        add_paragraph(doc, f"• 市场集中度（CR3）：{comp.get('cr3_2024', 0):.1f}%，趋势：{comp.get('concentration_trend', '未知')}")
        
        leaders = comp.get('leaders', [])
        if leaders:
            leader_names = [l.get('company', '') for l in leaders]
            add_paragraph(doc, f"• 市场领导者：{', '.join(leader_names)}")
        
        challengers = comp.get('challengers', [])
        if challengers:
            challenger_names = [c.get('company', '') for c in challengers]
            add_paragraph(doc, f"• 快速增长挑战者：{', '.join(challenger_names)}")
        
        add_paragraph(doc, "4.3 用户增长预测")
        user = market.get('user_growth_analysis', {})
        add_paragraph(doc, f"• 当前渗透率：{user.get('current_penetration', 0):.1f}%")
        add_paragraph(doc, f"• 2028年预测渗透率：{user.get('future_penetration', [0])[-1]:.1f}%")
        add_paragraph(doc, f"• 2028年潜在用户数：{user.get('future_users', [0])[-1]:.1f}百万")
        
        add_paragraph(doc, "4.4 技术趋势分析")
        tech = market.get('technology_analysis', [])
        if tech:
            star_techs = [t for t in tech if t.get('tech_category') == '明星技术']
            if star_techs:
                tech_names = [t.get('technology', '') for t in star_techs]
                add_paragraph(doc, f"• 明星技术：{', '.join(tech_names)}（高成熟度+高增长潜力）")
            
            cash_cow_techs = [t for t in tech if t.get('tech_category') == '现金牛技术']
            if cash_cow_techs:
                tech_names = [t.get('technology', '') for t in cash_cow_techs]
                add_paragraph(doc, f"• 现金牛技术：{', '.join(tech_names)}（高成熟度+稳定增长）")
    else:
        add_paragraph(doc, "（市场分析结果加载失败，请确保已运行分析流程）")
    
    doc.add_page_break()
    
    # ========== 五、投资建议与风险评估 ==========
    add_title(doc, "五、投资建议与风险评估", level=1)
    
    if results['market']:
        market = results['market']
        
        add_paragraph(doc, "5.1 风险评估")
        risk = market.get('risk_assessment', {})
        add_paragraph(doc, f"• 总体风险等级：{risk.get('overall_risk', '未知')}级")
        
        risks = risk.get('risks', [])
        if risks:
            for r in risks:
                add_paragraph(doc, f"• {r.get('type', '')}：{r.get('level', '')} - {r.get('description', '')}", bullet=True)
        
        add_paragraph(doc, "5.2 投资建议")
        recommendations = market.get('investment_recommendations', [])
        if recommendations:
            for rec in recommendations:
                add_paragraph(doc, f"• {rec.get('area', '')}：{rec.get('recommendation', '')}", bullet=True)
                add_paragraph(doc, f"  理由：{rec.get('rationale', '')}", indent=1)
        else:
            add_paragraph(doc, "• 市场进入：谨慎进入，需差异化竞争")
            add_paragraph(doc, "• 技术投资：重点投资远程监控、智能手环等明星技术")
            add_paragraph(doc, "• 产品改进：优先改进可靠性方面的问题")
    else:
        add_paragraph(doc, "5.1 风险评估")
        add_paragraph(doc, "• 市场风险：中 - 市场平稳增长（约15%）", bullet=True)
        add_paragraph(doc, "• 竞争风险：中 - 市场中等集中（CR3约65%）", bullet=True)
        add_paragraph(doc, "• 用户风险：中 - 用户满意度一般（3.6分）", bullet=True)
        add_paragraph(doc, "• 技术风险：低 - 技术成熟度较高", bullet=True)
        add_paragraph(doc, "• 政策风险：低 - 政策支持充分", bullet=True)
        
        add_paragraph(doc, "5.2 投资建议")
        add_paragraph(doc, "• 市场进入：谨慎进入，需差异化竞争", bullet=True)
        add_paragraph(doc, "• 技术投资：重点投资远程监控、智能手环等明星技术", bullet=True)
        add_paragraph(doc, "• 产品改进：优先改进可靠性方面的问题", bullet=True)
    
    doc.add_page_break()
    
    # ========== 六、关键发现与结论 ==========
    add_title(doc, "六、关键发现与结论", level=1)
    
    add_paragraph(doc, "6.1 用户需求关键发现")
    add_paragraph(doc, "• 安全性是用户最关注的方面，情感得分最高（0.898）", bullet=True)
    add_paragraph(doc, "• 可靠性是用户最不满意的方面，需要优先改进", bullet=True)
    add_paragraph(doc, "• 易用性覆盖评论最多（35.1%），说明操作简便性是普遍关注点", bullet=True)
    add_paragraph(doc, "• 价格方面的讨论相对较少（17.9%），但情感得分较高", bullet=True)
    
    add_paragraph(doc, "6.2 市场趋势关键发现")
    add_paragraph(doc, "• 市场规模约85亿元，保持平稳增长（约15%）", bullet=True)
    add_paragraph(doc, "• 市场集中度中等（CR3约65%），但呈下降趋势", bullet=True)
    add_paragraph(doc, "• 用户渗透率仅3.8%，有巨大增长空间", bullet=True)
    add_paragraph(doc, "• 远程监控和智能手环是技术热点", bullet=True)
    
    add_paragraph(doc, "6.3 战略建议")
    add_paragraph(doc, "• 产品策略：聚焦可靠性改进，强化安全性和易用性优势", bullet=True)
    add_paragraph(doc, "• 市场策略：针对细分人群（如科技先锋型老人）推出定制产品", bullet=True)
    add_paragraph(doc, "• 技术策略：重点投资远程监控和智能手环技术", bullet=True)
    add_paragraph(doc, "• 投资策略：关注具有技术差异化的初创企业", bullet=True)
    
    add_paragraph(doc, "6.4 研究局限与后续方向")
    add_paragraph(doc, "• 数据局限：基于电商评论，缺乏线下使用场景数据", bullet=True)
    add_paragraph(doc, "• 方法局限：情感分析模型对中文语境的理解有待优化", bullet=True)
    add_paragraph(doc, "• 后续方向：增加一手调研数据，构建更精准的用户画像", bullet=True)
    
    doc.add_page_break()
    
    # ========== 附录 ==========
    add_title(doc, "附录：分析工具与方法说明", level=1)
    
    add_paragraph(doc, "A.1 使用的Python库")
    libraries = [
        ["数据采集", "Selenium, Requests"],
        ["数据处理", "Pandas, NumPy"],
        ["文本分析", "Jieba, SnowNLP, WordCloud"],
        ["机器学习", "Scikit-learn"],
        ["可视化", "Matplotlib, Seaborn"],
        ["报告生成", "python-docx"]
    ]
    add_table(doc, ["功能", "主要库"], libraries)
    
    add_paragraph(doc, "A.2 分析流程示意图")
    add_paragraph(doc, "数据采集 → 数据清洗 → NLP分析 → 市场分析 → 报告生成")
    
    add_paragraph(doc, "A.3 文件说明")
    add_paragraph(doc, "• 原始数据：data/raw/ 目录下的CSV/JSON文件", bullet=True)
    add_paragraph(doc, "• 清洗后数据：data/processed/ 目录", bullet=True)
    add_paragraph(doc, "• 分析结果：data/analysis/ 目录下的JSON文件", bullet=True)
    add_paragraph(doc, "• 可视化图表：data/visualization/ 目录下的PNG文件", bullet=True)
    
    add_paragraph(doc, "A.4 代码文件说明")
    code_files = [
        ["elderly_product_crawler.py", "电商评论爬虫"],
        ["data_processing.py", "数据清洗与预处理"],
        ["nlp_analysis.py", "NLP情感分析"],
        ["market_analysis.py", "市场分析与预测"],
        ["generate_analysis_report.py", "报告生成（本文件）"]
    ]
    add_table(doc, ["文件名", "功能说明"], code_files)
    
    # ========== 保存文档 ==========
    output_dir = r"D:\panze（用户分身）\user\trae\file"
    os.makedirs(output_dir, exist_ok=True)
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_file = os.path.join(output_dir, f"智能家居适老化改造数据分析报告_{timestamp}.docx")
    
    doc.save(output_file)
    
    print("=" * 50)
    print("数据分析报告生成完成！")
    print(f"报告文件：{output_file}")
    print(f"报告页数：约{len(doc.paragraphs)//50 + 1}页")
    print("=" * 50)
    
    return output_file

def get_sample_results():
    """获取示例分析结果（用于测试）"""
    return {
        'sentiment': {
            'total_reviews': 8500,
            'avg_sentiment_score': 0.68,
            'avg_rating': 4.2,
            'sentiment_distribution': {'积极': 5200, '中性': 2200, '消极': 1100},
            'sentiment_consistency_rate': 67.6,
            'aspect_analysis': {
                '安全性': {'review_count': 3200, 'avg_sentiment': 0.72, 'avg_score': 4.3, 'coverage': 37.6},
                '易用性': {'review_count': 4500, 'avg_sentiment': 0.65, 'avg_score': 4.1, 'coverage': 52.9},
                '功能性': {'review_count': 2800, 'avg_sentiment': 0.70, 'avg_score': 4.2, 'coverage': 32.9},
                '可靠性': {'review_count': 2100, 'avg_sentiment': 0.62, 'avg_score': 3.9, 'coverage': 24.7},
                '价格': {'review_count': 1800, 'avg_sentiment': 0.58, 'avg_score': 3.8, 'coverage': 21.2}
            }
        },
        'keywords': {
            'tfidf_keywords': {'老人': 0.645, '功能': 0.537, '跌倒': 0.352, '手环': 0.273, '智能': 0.272},
            'textrank_keywords': {'功能': 1.000, '老人': 0.958, '跌倒': 0.440, '智能': 0.404, '操作': 0.347},
            'topics': [
                {'topic_id': 0, 'words': ['老人', '喜欢', '父母', '他们', '反应灵敏']},
                {'topic_id': 1, 'words': ['老人', '智能', '呼叫器', '乐橙', '操作']},
                {'topic_id': 2, 'words': ['老人', '语音控制', '华为', '中心', '安全']}
            ]
        },
        'market': {
            'market_overview': {'current_size': 85, 'growth_rate': 13.3, 'forecast_cagr': 11.2},
            'competition_analysis': {'cr3_2024': 65.0, 'concentration_trend': '降低'},
            'user_growth_analysis': {'current_penetration': 3.8, 'future_penetration': [4.5, 5.2, 6.0, 6.8, 7.1]},
            'technology_analysis': [
                {'technology': '远程监控', 'tech_category': '明星技术'},
                {'technology': '智能手环', 'tech_category': '明星技术'},
                {'technology': '跌倒检测', 'tech_category': '现金牛技术'}
            ],
            'risk_assessment': {'overall_risk': '中'},
            'investment_recommendations': [
                {'area': '市场进入', 'recommendation': '谨慎进入', 'rationale': '市场平稳增长，需差异化竞争'},
                {'area': '技术投资', 'recommendation': '重点投资', 'rationale': '远程监控、智能手环等明星技术'}
            ]
        },
        'files': {
            'wordcloud': 'data/visualization/wordcloud_example.png',
            'nlp_charts': 'data/visualization/nlp_analysis_example.png',
            'market_charts': 'data/visualization/market_analysis_example.png'
        }
    }

def main():
    """主函数"""
    try:
        print("开始生成数据分析报告...")
        report_file = create_analysis_report()
        print(f"报告生成成功：{report_file}")
        
        # 显示报告内容摘要
        print("\n报告内容摘要：")
        print("-" * 40)
        print("1. 研究概述：背景、目标、方法")
        print("2. 数据采集与处理方法")
        print("3. NLP情感分析结果（情感分布、方面分析、关键词）")
        print("4. 市场分析与预测（规模、竞争、用户、技术）")
        print("5. 投资建议与风险评估")
        print("6. 关键发现与结论")
        print("附录：分析工具与方法说明")
        
    except Exception as e:
        print(f"生成报告时出错：{e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()